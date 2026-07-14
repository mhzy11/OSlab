#!/usr/bin/env python3
"""Generate lab1 experiment report as docx."""

from docx import Document
from docx.shared import Pt, RGBColor
import re

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h

def add_code(doc, code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    doc.add_paragraph()

def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

# ===== Title =====
title = doc.add_heading('Lab1 实验报告：SW64 中断与异常处理', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ===== Section 1 =====
add_heading_styled(doc, '一、实验目的', 1)
purposes = [
    '了解 CPU 的中断机制，特别是 SW64 平台上 hmcode 对中断的中转处理',
    '掌握 SW64 上内核中断入口注册（wrent/wrkgp）的方法',
    '了解时钟中断的产生与控制（wrtimer oneshot 模式）',
    '掌握中断/异常处理中寄存器的保存与恢复（SAVE_ALL/RESTORE_ALL + HMC_rti）',
    '理解 CPU 异常机制（断点、非法指令），了解 hmcode 对不同异常类型的区分',
    '实现中断和异常处理的完整 C 语言分发框架'
]
for i, p in enumerate(purposes, 1):
    doc.add_paragraph(f'{i}. {p}')

# ===== Section 2 =====
add_heading_styled(doc, '二、实验环境', 1)
env_items = [
    '编译器：sw_64sw6b-sunway-linux-gnu-gcc（SW64 交叉编译器）',
    '模拟器：qemu-system-sw64（core3-hmcode 版本）',
    '机器模型：core3，单核，512M 内存',
    '内核入口地址：0xffffffff80910000',
    '内核打印区：[0xffffffff80700000, 0xffffffff80700000 + 1M)',
    'HMC 调用号：参考 Linux 内核 arch/sw_64/include/asm/hmcall.h'
]
for item in env_items:
    doc.add_paragraph(item, style='List Bullet')

# ===== Section 3 =====
add_heading_styled(doc, '三、实验原理', 1)

add_heading_styled(doc, '3.1 SW64 中断处理流程', 2)
doc.add_paragraph(
    'SW64 的中断处理与其他架构不同——硬件中断首先进入 hmcode（Hypervisor Microcode），'
    'hmcode 判断中断类型后跳转到内核注册的入口地址。流程如下：'
)
doc.add_paragraph(
    '硬件中断 → hmcode (host_interrupt1.S) → 读 CSR__INT_STAT 判断类型 → '
    '保存 PS/PC/GP/r16-r18 到内核栈 → 从 vcpucb 读取 entInt 地址 → '
    '跳转 entInt → SAVE_ALL → trap() → do_entInt()'
)

add_heading_styled(doc, '3.2 SW64 异常处理流程', 2)
doc.add_paragraph('异常（断点、非法指令）的入口是 entIF：')
doc.add_paragraph('断点：CPU 执行 sys_call 0x80 → hmcode 设置 r16=0 → entIF')
doc.add_paragraph('非法指令：CPU 解码 .long 0x7a000000 失败 → QEMU(打补丁) → hmcode opcdec.S 设置 r16=4 → entIF')
doc.add_paragraph('hmcode 通过 r16（a0）区分异常类型：0 = 断点，4 = 非法指令（opDEC）。')

add_heading_styled(doc, '3.3 内核入口注册机制', 2)
add_code(doc, '''void trap_init(void) {
    register unsigned long gptr __asm__("$29");
    wrkgp(gptr);                     // HMC 0x37: 写入内核 GP 到 vcpucb
    wrent((unsigned long)entInt, 0);  // HMC 0x34: 注册中断入口
    wrent((unsigned long)entIF, 3);   // HMC 0x34: 注册异常入口
}''')
doc.add_paragraph('wrent 的第二个参数指定入口类型：0 = 中断(entInt)，3 = 异常(entIF)。')

add_heading_styled(doc, '3.4 时钟中断（oneshot 模式）', 2)
doc.add_paragraph('SW64 的时钟是 oneshot 模式——触发一次后停止，必须在中断处理中重新 wrtimer。QEMU TCG 模式下定时器频率较低，wrtimer(1) 产生最短间隔。')

add_heading_styled(doc, '3.5 中断使能控制', 2)
doc.add_paragraph('swpipl 控制 CPU 的中断优先级（IPL）。IPL=0 允许所有中断，IPL>0 仅允许更高优先级的中断。')

add_heading_styled(doc, '3.6 HMC 调用接口', 2)
doc.add_paragraph('HMC 调用号来自 Linux 内核 hmcall.h，使用 sys_call 指令：')
add_table(doc,
    ['HMC', '编号', '功能'],
    [
        ['wrent', '0x34', '注册内核入口到 vcpucb'],
        ['swpipl', '0x35', '读写 CPU 中断优先级'],
        ['wrkgp', '0x37', '写入内核 GP（r29）到 vcpucb'],
        ['wrtimer', '0x3B', '设置 oneshot 定时器'],
        ['rti', '0x3F', '从异常/中断返回'],
        ['bpt', '0x80', '断点异常'],
    ]
)

# ===== Section 4 =====
add_heading_styled(doc, '四、实验步骤与代码修改', 1)

add_heading_styled(doc, '4.1 架构文件清理', 2)
doc.add_paragraph('从 lab0 继承 SW64 基础框架，删除 RISC-V 残留文件：libs/sbi.c、libs/sbi.h、libs/riscv.h。新建 libs/sw64.h 提供 SW64 汇编宏。')

add_heading_styled(doc, '4.2 Makefile 修改', 2)
doc.add_paragraph('新增 kern/debug、kern/driver、kern/trap、kern/mm 目录。关键修改：')
doc.add_paragraph('- -G 0 -fno-common：禁止 common symbol，避免 ticks 产生 .scommon 段导致 QEMU 加载失败')
doc.add_paragraph('- 链接 4 个预编译除法 .o 文件（SW64 无硬件除法）')

add_heading_styled(doc, '4.3 SW64 汇编入口实现（trapentry.S）', 2)
doc.add_paragraph('实现 entInt（中断入口）和 entIF（异常入口），SAVE_ALL 保存 r0~r28，RESTORE_ALL + HMC_rti 返回。hmcode 在跳转前已保存 PS/PC/GP/r16-r18 到栈上，SAVE_ALL 额外保存 r0-r28 和 cause，总计 36 槽位映射到 struct pushregs。')

add_heading_styled(doc, '4.4 寄存器保存布局（trap.h）', 2)
add_table(doc,
    ['偏移(槽)', '字段', '保存者', '说明'],
    [
        ['0-28', 'r0~r28', 'SAVE_ALL', '通用寄存器'],
        ['29', 'cause', 'entInt/entIF', '区分中断(0)/异常(3)'],
        ['30', 'PS', 'hmcode', '处理器状态'],
        ['31', 'PC', 'hmcode', '异常/中断地址'],
        ['32', 'GP', 'hmcode', '全局指针 r29'],
        ['33-35', 'r16~r18', 'hmcode', 'a0(类型)/a1/a2'],
    ]
)

add_heading_styled(doc, '4.5 C 语言中断/异常分发（trap.c）', 2)
doc.add_paragraph('trap() 根据 cause 分发：cause=0 → do_entInt（中断），cause≠0 → do_entIF（异常）。do_entInt 处理时钟中断（r16=9），do_entIF 处理断点（r16=0）和非法指令（r16=4）。')
doc.add_paragraph('关键细节：断点的 hmcode 保存的是返回地址（不额外 +4），非法指令的 QEMU 补丁保存的是异常指令地址（需 +4 跳过）。')

add_heading_styled(doc, '4.6 驱动层实现', 2)
doc.add_paragraph('clock.c：wrtimer 封装；intr.c：swpipl 封装；sw.h：所有 HMC 调用 inline 函数，参考 Linux __CALL_HMC_W1/__CALL_HMC_W2 宏的 clobber 约定。')

add_heading_styled(doc, '4.7 QEMU 非法指令补丁', 2)
add_code(doc, '''--- a/target/sw64/helper.c
+++ b/target/sw64/helper.c
     case EXCP_OPCDEC:
-        cpu_abort(cs, "ILLEGAL INSN");
+        i = 0xD80;
+        env->csr[EXC_PC] = env->pc - 4;
         break;''')

add_heading_styled(doc, '4.8 init.c 集成测试', 2)
add_code(doc, '''int kern_init(void) {
    memset(edata, 0, end - edata);
    cprintf("(THU.CST) os is loading ...\\n\\n");
    trap_init();                         // 注册 entInt/entIF
    clock_init();                        // 启动 oneshot 定时器
    intr_enable();                       // 开中断
    asm volatile("sys_call 0x80"::);     // 触发断点异常
    asm volatile(".long 0x7a000000"::);  // 触发非法指令异常
    while (1);                           // 等待时钟中断
}''')

# ===== Section 5 =====
add_heading_styled(doc, '五、遇到的问题与解决方案', 1)

add_heading_styled(doc, '问题 1：.scommon 段导致 QEMU 内核加载失败', 2)
doc.add_paragraph('现象：make 编译通过，QEMU 启动后无输出，pmemsave 无 p.log。')
doc.add_paragraph('定位：对比 lab0 和 lab1 的 ELF program headers，发现 lab1 的 LOAD segment 多了 .scommon 段（MemSiz 多 8 字节）。该段由 clock.c 和 trap.c 中 volatile size_t ticks 的 tentative definition 产生。')
doc.add_paragraph('解决：Makefile 加 -fno-common，同时将 trap.c 中重复的 ticks 定义改为 extern，只保留 clock.c 中一份定义。')

add_heading_styled(doc, '问题 2：断点处理导致非法指令被跳过', 2)
doc.add_paragraph('现象：p.log 中只有 breakpoint pc，没有 opDEC pc。')
doc.add_paragraph('定位：do_entIF 的 case 0（断点）中多加了 tf->gpr.pc += 4。但 hmcode 对 sys_call 0x80 保存的是返回地址（已在指令之后），加 4 后把紧随其后的非法指令也跳过去了。')
doc.add_paragraph('解决：断点 case 去掉 tf->gpr.pc += 4，仅保留非法指令 case。')

add_heading_styled(doc, '问题 3：非法指令 PC 显示错误', 2)
doc.add_paragraph('现象：opDEC pc 与 breakpoint 同地址。')
doc.add_paragraph('定位：非法指令 case 打印 tf->gpr.pc - 4，但 QEMU 补丁已把 PC 设为 env->pc - 4（即异常指令本身地址），再减 4 多减了一次。')
doc.add_paragraph('解决：非法指令 case 改为打印 tf->gpr.pc（不下减 4），正确显示比断点地址多 4 字节。')

# ===== Section 6 =====
add_heading_styled(doc, '六、运行验证', 1)

add_heading_styled(doc, '编译', 2)
add_code(doc, 'make clean && make')

add_heading_styled(doc, '运行', 2)
add_code(doc, 'make qemu')

add_heading_styled(doc, '验证输出', 2)
add_code(doc, '# 在 QEMU monitor 中：\npmemsave 0x700000 1000 p.log\nquit\ncat p.log')

add_heading_styled(doc, '验证结果', 2)
add_code(doc, '''(THU.CST) os is loading ...

++ setup timer interrupts
breakpoint pc = ffffffff809100c8
opDEC  pc = ffffffff809100cc
100 ticks
200 ticks
300 ticks
400 ticks
...''')

add_table(doc,
    ['输出', '验证内容'],
    [
        ['(THU.CST) os is loading ...', '内核正常启动'],
        ['++ setup timer interrupts', '时钟初始化成功'],
        ['breakpoint pc = ffffffff809100c8', '断点异常捕获（sys_call 0x80）'],
        ['opDEC pc = ffffffff809100cc', '非法指令异常捕获（.long 0x7a000000）'],
        ['100 ticks → 200 ticks → ...', '时钟中断持续触发，oneshot 重设正常'],
    ]
)
doc.add_paragraph('断点和非法指令地址相差 4 字节（紧邻的两条指令），中断计数持续递增，证明寄存器保存/恢复正确无误。')

# ===== Section 7 =====
add_heading_styled(doc, '七、实验总结', 1)
doc.add_paragraph('本次实验在 lab0 SW64 最小内核基础上，实现了完整的中断和异常处理框架：')

summary_items = [
    'HMC 接口封装：实现 wrent/wrkgp/wrtimer/swpipl/rti 的 inline assembly 封装，调用号来自 Linux 内核 hmcall.h',
    '汇编入口：实现 entInt/entIF 入口，SAVE_ALL 保存寄存器，__trapret + HMC_rti 返回',
    '中断分发：trap() 根据 cause 字段分发中断/异常，do_entInt 处理时钟中断，do_entIF 处理断点和非法指令',
    '时钟驱动：wrtimer oneshot 模式，每次中断后重新触发',
    '问题排查：.scommon 段 QEMU 兼容性、断点/非法指令的 PC 处理差异'
]
for item in summary_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    '核心理解：SW64 的中断机制与 RISC-V 本质不同——硬件不直接跳转内核入口，而是通过 hmcode 中转。'
    'hmcode 负责保存部分上下文、区分中断类型，内核通过 wrent 注册入口、wrkgp 注册 GP 来建立与 hmcode 的约定。'
)

# Save
doc.save('实验报告_lab1.docx')
print('Done: 实验报告_lab1.docx')
